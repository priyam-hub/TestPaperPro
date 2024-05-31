import streamlit as st
from langchain_groq import ChatGroq
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
import os
import docx
from io import BytesIO  


llm = ChatGroq(temperature=0.5, groq_api_key="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", model_name="llama3-8b-8192")


# Placeholder functions to simulate LLM invocation (replace with actual LLM code)
def generate_mcq_questions(subject_name, syllabus, num_mcq, example_mcq):
    prompt_template = """
    You are an expert in creating educational content. Based on the following inputs, generate {num_mcq} multiple-choice questions (MCQs) for a quiz.

    Subject: {subject_name}
    Syllabus: {syllabus}
    Example MCQ question format:
    {example_mcq}

    Please choose the correct alternatives for the following questions.
    Please do not provide the answers to the questions.

    Q1. _________________
    (a)..... (b)..... (c)..... (d)......
    
    Q2. _________________
    (a)..... (b)..... (c)..... (d)......
    
    ...
    
    Generate the questions in the same format.
    """

    prompt = prompt_template.format(
        subject_name=subject_name,
        syllabus=syllabus,
        num_mcq=num_mcq,
        example_mcq=example_mcq
    )

    chain = (
        ChatPromptTemplate.from_template(prompt)
        | llm
        | StrOutputParser()
    )

    return chain.invoke({})

def generate_short_questions(subject_name, syllabus, num_short, example_short):
    prompt_template = """
    You are an expert in creating educational content. Based on the following inputs, generate {num_short} short answer questions for a quiz.

    Subject: {subject_name}
    Syllabus: {syllabus}
    Example short question format:
    {example_short}

    Please do not provide the answers to the questions.
    Do not provide any notes to the questions.

    Q1. _________________
    
    Q2. _________________
    
    ...
    
    Generate the questions in the same format.
    """

    prompt = prompt_template.format(
        subject_name=subject_name,
        syllabus=syllabus,
        num_short=num_short,
        example_short=example_short
    )

    chain = (
        ChatPromptTemplate.from_template(prompt)
        | llm
        | StrOutputParser()
    )

    return chain.invoke({})

def generate_long_questions(subject_name, syllabus, num_long, example_long):
    prompt_template = """
    You are an expert in creating educational content. Based on the following inputs, generate {num_long} long answer questions for a quiz.

    Subject: {subject_name}
    Syllabus: {syllabus}
    Example long question format:
    {example_long}

    Please do not provide the answers to the questions.
    Do not provide any notes to the questions.

    Q1. _________________
    
    Q2. _________________
    
    ...
    
    Generate the questions in the same format.
    """

    prompt = prompt_template.format(
        subject_name=subject_name,
        syllabus=syllabus,
        num_long=num_long,
        example_long=example_long
    )

    chain = (
        ChatPromptTemplate.from_template(prompt)
        | llm
        | StrOutputParser()
    )

    return chain.invoke({})

# Initialize Streamlit session state
if "mcq_questions" not in st.session_state:
    st.session_state.mcq_questions = ""
if "short_questions" not in st.session_state:
    st.session_state.short_questions = ""
if "long_questions" not in st.session_state:
    st.session_state.long_questions = ""

# Streamlit app code
st.title("Dynamic Test Paper Generator using Bloom's Taxonomy")

#Sidebar Logo
st.sidebar.image("Logo.png")

# Sidebar inputs
st.sidebar.header("Input Details")

# Subject Name
subject_name = st.sidebar.text_input("Subject Name")
if subject_name:
    st.sidebar.markdown("✅ Subject Name entered")

# Syllabus
syllabus = st.sidebar.text_area("Syllabus")
if syllabus:
    st.sidebar.markdown("✅ Syllabus entered")

# Number of MCQ questions
num_mcq = st.sidebar.slider("Number of MCQ questions", 0, 100, 10)
st.sidebar.markdown(f"✅ Number of MCQ questions: {num_mcq}")

# Example MCQ question
example_mcq = st.sidebar.text_area("Example MCQ question")
if example_mcq:
    st.sidebar.markdown("✅ Example MCQ question entered")

# Number of short questions
num_short = st.sidebar.slider("Number of short questions", 0, 100, 5)
st.sidebar.markdown(f"✅ Number of short questions: {num_short}")

# Example short question
example_short = st.sidebar.text_area("Example short question")
if example_short:
    st.sidebar.markdown("✅ Example short question entered")

# Number of long answer questions
num_long = st.sidebar.slider("Number of long answer questions", 0, 50, 3)
st.sidebar.markdown(f"✅ Number of long answer questions: {num_long}")

# Example long answer question
example_long = st.sidebar.text_area("Example long answer question")
if example_long:
    st.sidebar.markdown("✅ Example long answer question entered")
    

# Button to generate MCQ questions
if st.sidebar.button("Generate MCQ Questions"):
    if subject_name and syllabus and example_mcq:
        st.session_state.mcq_questions = generate_mcq_questions(subject_name, syllabus, num_mcq, example_mcq)
        st.header("Generated MCQ Questions")
        st.text_area("Generated MCQ Questions", value=st.session_state.mcq_questions, height=400)
    else:
        st.sidebar.markdown("❗ Please fill in all required fields to generate MCQ questions")

# Button to generate Short Questions
if st.sidebar.button("Generate Short Questions"):
    if subject_name and syllabus and example_short:
        st.session_state.short_questions = generate_short_questions(subject_name, syllabus, num_short, example_short)
        st.header("Generated Short Questions")
        st.text_area("Generated Short Questions", value=st.session_state.short_questions, height=400)
    else:
        st.sidebar.markdown("❗ Please fill in all required fields to generate short answer questions")

# Button to generate Long Questions
if st.sidebar.button("Generate Long Questions"):
    if subject_name and syllabus and example_long:
        st.session_state.long_questions = generate_long_questions(subject_name, syllabus, num_long, example_long)
        st.header("Generated Long Questions")
        st.text_area("Generated Long Questions", value=st.session_state.long_questions, height=400)
    else:
        st.sidebar.markdown("❗ Please fill in all required fields to generate long answer questions")
        
# Display all questions before downloading
if st.session_state.mcq_questions or st.session_state.short_questions or st.session_state.long_questions:
    st.header("All Generated Questions")
    all_questions = ""
    if st.session_state.mcq_questions:
        all_questions += "Generated MCQ Questions:\n" + st.session_state.mcq_questions + "\n\n"
    if st.session_state.short_questions:
        all_questions += "Generated Short Answer Questions:\n" + st.session_state.short_questions + "\n\n"
    if st.session_state.long_questions:
        all_questions += "Generated Long Answer Questions:\n" + st.session_state.long_questions + "\n\n"
    st.text_area("All Generated Questions", value=all_questions, height=600)
        
# Button to download questions as a DOCX file
if st.sidebar.button("Generate All Questions as DOCX"):
    if st.session_state.mcq_questions or st.session_state.short_questions or st.session_state.long_questions:
        doc = docx.Document()
        if st.session_state.mcq_questions:
            doc.add_heading("Generated MCQ Questions", level=1)
            doc.add_paragraph(st.session_state.mcq_questions)
        if st.session_state.short_questions:
            doc.add_heading("Generated Short Answer Questions", level=1)
            doc.add_paragraph(st.session_state.short_questions)
        if st.session_state.long_questions:
            doc.add_heading("Generated Long Answer Questions", level=1)
            doc.add_paragraph(st.session_state.long_questions)

        # Save the document to a BytesIO object
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Offer the file for download
        st.download_button(
            label="Download Generated Questions as DOCX",
            data=doc_io,
            file_name="generated_questions.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.sidebar.markdown("❗ Generate some questions first before downloading")